# Copyright (c) 2024 IBM Corp. All rights reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""Classes for populate FedRAMP Docx template."""

import logging
from typing import Dict, List, Optional, Tuple

from docx.document import Document  # type: ignore
from docx.table import Table, _Cell  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

from trestle.common.err import TrestleError

import trestle_fedramp.const as const
from trestle_fedramp.core.ssp_reader import FedrampControlDict, FedrampSSPData

logger = logging.getLogger(__name__)


class FedrampDocx():
    """Class for populating the FedRAMP template."""

    def __init__(self, docx: Document, control_dict: FedrampControlDict) -> None:
        """Initialize the FedrampDocx class."""
        self.docx = docx
        self.control_dict = control_dict

        self.row_header_location: Tuple[int, int] = (0, 0)

        self.control_summaries = ControlSummaries()
        self.control_implementation_descriptions = ControlImplementationDescriptions()

    def populate(self) -> None:
        """Populate the FedRAMP template."""
        try:
            for table in self.docx.tables:
                row_header: str = table.cell(*self.row_header_location).text
                if self.control_summaries.is_control_summary_table(row_header):
                    control_id = self.get_control_id(row_header)
                    ssp_data: Optional[FedrampSSPData] = self.get_control_data(control_id)
                    if ssp_data:
                        self.control_summaries.populate_table(table, control_id, ssp_data)

                elif self.control_implementation_descriptions.is_control_implementation_table(row_header):
                    control_id = self.get_control_id(row_header)
                    ssp_data = self.get_control_data(control_id)
                    if ssp_data:
                        self.control_implementation_descriptions.populate_table(table, control_id, ssp_data)
        except Exception as e:
            raise TrestleError(f'Error populating FedRAMP template: {e}')

    def get_control_data(self, control_id: str) -> Optional[FedrampSSPData]:
        """Get the control data from the control_dict."""
        if control_id in self.control_dict:
            logging.debug(f'Found the control info for {control_id}.')
            ssp_data: FedrampSSPData = self.control_dict[control_id]
            return ssp_data
        else:
            logging.debug(f'Control {control_id} not found in the SSP data.')
            return None

    @staticmethod
    def get_control_id(row_header: str) -> str:
        """Get the control id from the table."""
        return row_header.split(' ')[0]


class ControlSummaries():
    """Populate the control summaries in the FedRAMP template."""

    def __init__(self) -> None:
        """Initialize the ControlSummaries class."""
        # Control origination is always the last row in the table
        self._control_origination_cell: Tuple[int, int] = (-1, 0)
        self._control_origination_index_values: Dict[str, int] = {
            const.FEDRAMP_SP_CORPORATE: 1,
            const.FEDRAMP_SP_SYSTEM: 2,
            const.FEDRAMP_HYBRID: 3,
            const.FEDRAMP_CUST_CONFIGURED: 4,
            const.FEDRAMP_CUST_PROVIDED: 5,
            const.FEDRAMP_SHARED: 6,
            const.FEDRAMP_INHERITED: 7
        }
        # Implementation status is always the second last row in the table
        self._implementation_status_cell: Tuple[int, int] = (-2, 0)
        self._implementation_status_index_values: Dict[str, int] = {
            const.FEDRAMP_IMPLEMENTED: 1,
            const.FEDRAMP_PARTIAL: 2,
            const.FEDRAMP_PLANNED: 3,
            const.FEDRAMP_ALTERNATIVE: 4,
            const.FEDRAMP_NOT_APPLICABLE: 5
        }
        self._parameter_start_row = 2
        self._responsible_role_cell: Tuple[int, int] = (1, 0)

    @staticmethod
    def is_control_summary_table(row_header: str) -> bool:
        """Check if the table is a control summary table."""
        return row_header.endswith(const.CONTROL_SUMMARY)

    def get_control_origination_index(self, control_origination: str) -> int:
        """
        Get paragraph index in the control origination cell by control origination string.

        Args:
            control_origination: The control origination string.

        Returns:
            The paragraph index location in the control origination cell
            in the control summary table.
        """
        if control_origination not in self._control_origination_index_values:
            raise TrestleError(f'Invalid FedRAMP control origination value: {control_origination}')
        return self._control_origination_index_values[control_origination]

    def get_implementation_status_index(self, implementation_status: str) -> int:
        """
        Get paragraph index in the implementation status cell by implementation status string.

        Args:
            implementation_status: The implementation status string.

        Returns:
            The paragraph index location in the implementation status cell
            in the control summary table.
        """
        if implementation_status not in self._implementation_status_index_values:
            raise TrestleError(f'Invalid FedRAMP implementation status value: {implementation_status}')
        return self._implementation_status_index_values[implementation_status]

    def _set_checkbox(self, paragraph: Paragraph) -> None:
        """Check the checkbox in the paragraph."""
        # Find the checkbox element and set the checked attribute to 1
        check_box = paragraph._element.xpath(const.CHECKBOX_XPATH)[0]
        if check_box is None:
            raise TrestleError(f'Checkbox not found in the paragraph with text: {paragraph.text}')
        checked = check_box.find(f'{const.XML_NAMESPACE}checked')
        checked.attrib[f'{const.XML_NAMESPACE}val'] = '1'
        self._set_checkbox_text(paragraph)

    def _set_checkbox_text(self, paragraph: Paragraph) -> None:
        """Set the checkbox text."""
        checkbox_text = paragraph._element.xpath(const.BOX_ICON_XPATH)[0]
        if checkbox_text is None:
            raise TrestleError(f'Checkbox text not found in the paragraph with text: {paragraph.text}')
        logger.debug(f'Checkbox text found in the paragraph with text: {paragraph.text}')
        checkbox_text.text = const.CHECKED_BOX_ICON

    def _set_control_origination(self, control_origination_cell: _Cell, control_origination_values: List[str]) -> None:
        """Set the control origination in the cell."""
        for control_origination in control_origination_values:
            co_paragraph_index_loc = self.get_control_origination_index(control_origination)
            # Control origination is always the last row in the table
            if co_paragraph_index_loc > len(control_origination_cell.paragraphs):
                raise TrestleError(f'Invalid control origination: {control_origination}')
            co_paragraph: Paragraph = control_origination_cell.paragraphs[co_paragraph_index_loc]
            self._set_checkbox(co_paragraph)

    def _set_implementation_status(self, implementation_status_cell: _Cell, implementation_status: str) -> None:
        """Set the implementation status in the cell."""
        is_paragraph_index_loc = self.get_implementation_status_index(implementation_status)
        is_paragraph: Paragraph = implementation_status_cell.paragraphs[is_paragraph_index_loc]
        if is_paragraph_index_loc > len(implementation_status_cell.paragraphs):
            raise TrestleError(f'Invalid implementation status: {implementation_status}')
        self._set_checkbox(is_paragraph)

    def _set_parameter_values(self, table: Table, parameter_values: Dict[str, str]) -> None:
        """
        Set the parameter values in the cell.

        Args:
            parameter_cells: The list of parameter cells.
            parameter_values: The parameter values.

        Notes: In the control summary table parameters start in the second row after responsible role and
        continue for a non-standard number of rows until the implementation status is reached. All other values
        of the control summary stable have a standard location and number of rows.
        """
        for cell in table.columns[0].cells[self._parameter_start_row:]:
            label = self.get_parameter_id(cell.text)
            parameter_text = parameter_values.get(label, '')
            if parameter_text:
                cell.text = cell.text + ' ' + parameter_text

    @staticmethod
    def get_parameter_id(row_text: str) -> str:
        """Get the parameter id from the row text."""
        if row_text.startswith(const.FEDRAMP_PARAMETER_PREFIX):
            return row_text.split(' ')[1].strip(':')
        else:
            return ''

    def populate_table(self, table: Table, control_id: str, ssp_data: FedrampSSPData) -> None:
        """Populate the table with the SSP data."""
        try:
            if ssp_data.control_origination:
                control_origination_cell: _Cell = table.cell(*self._control_origination_cell)
                self._set_control_origination(control_origination_cell, ssp_data.control_origination)
            if ssp_data.implementation_status:
                implementation_status_cell: _Cell = table.cell(*self._implementation_status_cell)
                self._set_implementation_status(implementation_status_cell, ssp_data.implementation_status)
            if ssp_data.parameters:
                self._set_parameter_values(table, ssp_data.parameters)
            if ssp_data.responsible_roles:
                responsible_role_cell: _Cell = table.cell(*self._responsible_role_cell)
                responsible_role_str = ', '.join(ssp_data.responsible_roles)
                responsible_role_cell.text = f'{responsible_role_cell.text} {responsible_role_str}'
        except Exception as e:
            raise TrestleError(f'Error populating control summary for {control_id}: {e}')


class ControlImplementationDescriptions():
    """Populate the control implementation description in the FedRAMP template."""

    @staticmethod
    def is_control_implementation_table(row_header: str) -> bool:
        """Check if the table is a control implementation table."""
        return row_header.endswith(const.CONTROL_RESPONSE)

    @staticmethod
    def get_part_id(part_information: str) -> str:
        """Get the part id from the table."""
        if part_information.startswith(const.FEDRAMP_STATEMENT_PREFIX):
            return part_information.split(' ')[1].strip(':')
        else:
            return ''

    def populate_table(self, table: Table, control_id: str, ssp_data: FedrampSSPData) -> None:
        """Populate the table with the SSP data."""
        if ssp_data.control_implementation_description:
            try:
                # After the row header for each control implementation
                for cell in table.columns[0].cells[1:]:
                    label = self.get_part_id(cell.text)
                    if label in ssp_data.control_implementation_description:
                        description_text: str = ssp_data.control_implementation_description[label]
                        paragraph: Paragraph = cell.add_paragraph(description_text)
                        paragraph.add_run().add_break()
            except TrestleError as e:
                raise TrestleError(f'Error populating control implementation description for {control_id}: {e}')
